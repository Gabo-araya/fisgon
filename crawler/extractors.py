# crawler/extractors.py
import os
import logging
from typing import Dict, Any, Optional
from datetime import datetime
import hashlib

logger = logging.getLogger('crawler')

# Librerías para extracción de metadatos
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 no disponible. Instalar con: pip install PyPDF2")

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False
    logger.warning("Pillow no disponible. Instalar con: pip install Pillow")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx no disponible. Instalar con: pip install python-docx")

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl no disponible. Instalar con: pip install openpyxl")

try:
    from mutagen import File as MutagenFile
    from mutagen.mp3 import MP3
    from mutagen.mp4 import MP4
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False
    logger.warning("mutagen no disponible. Instalar con: pip install mutagen")

try:
    from odf.opendocument import load
    from odf import text, table, draw
    from odf.element import Text
    ODFPY_AVAILABLE = True
except ImportError:
    ODFPY_AVAILABLE = False
    logger.warning("odfpy no disponible. Instalar con: pip install odfpy")

try:
    import json
    import xml.etree.ElementTree as ET
    XML_JSON_AVAILABLE = True
except ImportError:
    XML_JSON_AVAILABLE = False


class MetadataExtractor:
    '''Clase base para extractores de metadatos'''
    
    def __init__(self, file_path: str, file_url: str = None, referrer: str = None):
        self.file_path = file_path
        self.file_url = file_url
        self.referrer = referrer
        self.metadata = {}
    
    def extract(self) -> Dict[str, Any]:
        '''Método principal que debe ser implementado por cada extractor'''
        raise NotImplementedError
    
    def get_common_metadata(self) -> Dict[str, Any]:
        '''Obtiene metadatos comunes a todos los archivos'''
        try:
            stat_info = os.stat(self.file_path)
            file_hash = self._calculate_file_hash()
            
            return {
                'file_path': self.file_path,
                'file_url': self.file_url,
                'referrer': self.referrer,
                'file_size': stat_info.st_size,
                'file_hash_sha256': file_hash,
                'created_at': datetime.fromtimestamp(stat_info.st_ctime).isoformat(),
                'modified_at': datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
                'extracted_at': datetime.now().isoformat(),
            }
        except Exception as e:
            logger.error(f"Error obteniendo metadatos comunes: {str(e)}")
            return {}
    
    def _calculate_file_hash(self) -> str:
        '''Calcula el hash SHA-256 del archivo'''
        try:
            with open(self.file_path, 'rb') as f:
                return hashlib.sha256(f.read()).hexdigest()
        except Exception:
            return ""


class PDFExtractor(MetadataExtractor):
    '''Extractor especializado para archivos PDF'''
    
    def extract(self) -> Dict[str, Any]:
        metadata = self.get_common_metadata()
        
        if not PDF_AVAILABLE:
            logger.warning("PyPDF2 no disponible, extracción limitada")
            return metadata
        
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Metadatos del documento
                if pdf_reader.metadata:
                    pdf_metadata = {
                        'author': self._clean_text(pdf_reader.metadata.get('/Author', '')),
                        'creator': self._clean_text(pdf_reader.metadata.get('/Creator', '')),
                        'producer': self._clean_text(pdf_reader.metadata.get('/Producer', '')),
                        'title': self._clean_text(pdf_reader.metadata.get('/Title', '')),
                        'subject': self._clean_text(pdf_reader.metadata.get('/Subject', '')),
                        'keywords': self._clean_text(pdf_reader.metadata.get('/Keywords', '')),
                        'creation_date': self._parse_pdf_date(pdf_reader.metadata.get('/CreationDate')),
                        'modification_date': self._parse_pdf_date(pdf_reader.metadata.get('/ModDate')),
                    }
                    
                    # Filtrar valores vacíos
                    pdf_metadata = {k: v for k, v in pdf_metadata.items() if v}
                    metadata['pdf_metadata'] = pdf_metadata
                
                # Información del PDF
                metadata['pdf_info'] = {
                    'num_pages': len(pdf_reader.pages),
                    'encrypted': pdf_reader.is_encrypted,
                    'pdf_version': getattr(pdf_reader, 'pdf_header', '').replace('%PDF-', '') if hasattr(pdf_reader, 'pdf_header') else '',
                }
                
                # Extraer texto de la primera página para análisis
                if len(pdf_reader.pages) > 0:
                    first_page_text = pdf_reader.pages[0].extract_text()[:500]
                    metadata['first_page_preview'] = first_page_text
                
        except Exception as e:
            logger.error(f"Error extrayendo metadatos de PDF {self.file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def extract_full_content(self) -> str:
        '''Extrae texto completo de todas las páginas del PDF'''
        if not PDF_AVAILABLE:
            logger.warning("PyPDF2 no disponible para extracción de contenido")
            return ""
        
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                full_text = []
                
                # Limitar a máximo 500 páginas por performance
                max_pages = min(len(pdf_reader.pages), 500)
                logger.info(f"Extrayendo contenido de {max_pages} páginas de PDF")
                
                for i, page in enumerate(pdf_reader.pages[:max_pages]):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            full_text.append(f"=== Página {i+1} ===\n{page_text.strip()}")
                    except Exception as e:
                        logger.warning(f"Error extrayendo página {i+1}: {str(e)}")
                        continue
                
                result = '\n\n'.join(full_text)
                logger.info(f"PDF contenido extraído: {len(result)} caracteres")
                return result
                
        except Exception as e:
            logger.error(f"Error extrayendo contenido completo PDF: {str(e)}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        '''Limpia texto de caracteres especiales de PDF'''
        if not text:
            return ""
        return text.strip().replace('\x00', '').replace('\r', '').replace('\n', ' ')
    
    def _parse_pdf_date(self, date_str: str) -> str:
        '''Parsea fechas en formato PDF a ISO'''
        if not date_str:
            return ""
        
        try:
            # Formato típico: D:20230515143022+05'00'
            if date_str.startswith('D:'):
                date_part = date_str[2:16]  # YYYYMMDDHHMMSS
                if len(date_part) >= 8:
                    year = date_part[:4]
                    month = date_part[4:6]
                    day = date_part[6:8]
                    hour = date_part[8:10] if len(date_part) >= 10 else "00"
                    minute = date_part[10:12] if len(date_part) >= 12 else "00"
                    second = date_part[12:14] if len(date_part) >= 14 else "00"
                    
                    return f"{year}-{month}-{day}T{hour}:{minute}:{second}"
        except Exception:
            pass
        
        return str(date_str)


class ImageExtractor(MetadataExtractor):
    '''Extractor especializado para imágenes (JPEG, PNG, TIFF, etc.)'''
    
    def extract(self) -> Dict[str, Any]:
        metadata = self.get_common_metadata()
        
        if not PILLOW_AVAILABLE:
            logger.warning("Pillow no disponible, extracción limitada")
            return metadata
        
        try:
            with Image.open(self.file_path) as image:
                # Información básica de la imagen
                metadata['image_info'] = {
                    'format': image.format,
                    'mode': image.mode,
                    'size': image.size,
                    'width': image.width,
                    'height': image.height,
                }
                
                # Extraer EXIF si está disponible
                if hasattr(image, '_getexif') and image._getexif():
                    exif_data = image._getexif()
                    if exif_data:
                        exif_metadata = {}
                        
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)
                            
                            # Procesar tags importantes
                            if tag == 'DateTime':
                                exif_metadata['datetime_original'] = str(value)
                            elif tag == 'DateTimeOriginal':
                                exif_metadata['datetime_original'] = str(value)
                            elif tag == 'Make':
                                exif_metadata['camera_make'] = str(value)
                            elif tag == 'Model':
                                exif_metadata['camera_model'] = str(value)
                            elif tag == 'Software':
                                exif_metadata['software'] = str(value)
                            elif tag == 'Artist':
                                exif_metadata['artist'] = str(value)
                            elif tag == 'Copyright':
                                exif_metadata['copyright'] = str(value)
                            elif tag == 'GPSInfo':
                                gps_data = self._extract_gps_data(value)
                                if gps_data:
                                    exif_metadata['gps_coordinates'] = gps_data
                        
                        if exif_metadata:
                            metadata['exif_metadata'] = exif_metadata
                
        except Exception as e:
            logger.error(f"Error extrayendo metadatos de imagen {self.file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _extract_gps_data(self, gps_info: dict) -> Optional[Dict[str, Any]]:
        '''Extrae y convierte coordenadas GPS'''
        try:
            gps_data = {}
            
            for tag_id, value in gps_info.items():
                tag = GPSTAGS.get(tag_id, tag_id)
                gps_data[tag] = value
            
            # Convertir coordenadas a decimal
            if 'GPSLatitude' in gps_data and 'GPSLongitude' in gps_data:
                lat = self._convert_to_degrees(gps_data['GPSLatitude'])
                lon = self._convert_to_degrees(gps_data['GPSLongitude'])
                
                # Aplicar signos según hemisferio
                if gps_data.get('GPSLatitudeRef') == 'S':
                    lat = -lat
                if gps_data.get('GPSLongitudeRef') == 'W':
                    lon = -lon
                
                return {
                    'latitude': lat,
                    'longitude': lon,
                    'coordinates_string': f"({lat:.6f}, {lon:.6f})"
                }
        
        except Exception as e:
            logger.error(f"Error procesando GPS: {str(e)}")
        
        return None
    
    def _convert_to_degrees(self, value):
        '''Convierte coordenadas GPS de grados/minutos/segundos a decimal'''
        d, m, s = value
        return float(d) + float(m)/60.0 + float(s)/3600.0


class OfficeExtractor(MetadataExtractor):
    '''Extractor para documentos de Microsoft Office (Word, Excel, PowerPoint)'''
    
    def extract(self) -> Dict[str, Any]:
        metadata = self.get_common_metadata()
        
        file_extension = os.path.splitext(self.file_path)[1].lower()
        
        if file_extension in ['.docx']:
            return self._extract_docx_metadata(metadata)
        elif file_extension in ['.xlsx']:
            return self._extract_xlsx_metadata(metadata)
        elif file_extension in ['.xls']:
            return self._extract_xls_metadata(metadata)
        else:
            # Para otros formatos de Office (.doc, .ppt)
            metadata['office_format'] = file_extension
            metadata['extraction_note'] = 'Formato de Office legacy, extracción limitada'
        
        return metadata
    
    def _extract_docx_metadata(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        '''Extrae metadatos de documentos Word (.docx)'''
        if not DOCX_AVAILABLE:
            metadata['extraction_error'] = 'python-docx no disponible'
            return metadata
        
        try:
            doc = Document(self.file_path)
            core_props = doc.core_properties
            
            office_metadata = {}
            
            # Propiedades principales
            if core_props.author:
                office_metadata['author'] = core_props.author
            if core_props.last_modified_by:
                office_metadata['last_modified_by'] = core_props.last_modified_by
            if core_props.created:
                office_metadata['created'] = core_props.created.isoformat()
            if core_props.modified:
                office_metadata['modified'] = core_props.modified.isoformat()
            if core_props.title:
                office_metadata['title'] = core_props.title
            if core_props.subject:
                office_metadata['subject'] = core_props.subject
            if core_props.keywords:
                office_metadata['keywords'] = core_props.keywords
            if core_props.comments:
                office_metadata['comments'] = core_props.comments
            if core_props.category:
                office_metadata['category'] = core_props.category
            if core_props.revision:
                office_metadata['revision'] = core_props.revision
            
            metadata['office_metadata'] = office_metadata
            
            # Información del documento
            metadata['document_info'] = {
                'paragraphs_count': len(doc.paragraphs),
                'document_type': 'Word Document (.docx)'
            }
            
            # Extraer texto inicial para análisis
            first_paragraphs = []
            for para in doc.paragraphs[:3]:
                if para.text.strip():
                    first_paragraphs.append(para.text.strip())
            
            if first_paragraphs:
                metadata['content_preview'] = ' '.join(first_paragraphs)[:500]
            
        except Exception as e:
            logger.error(f"Error extrayendo metadatos de DOCX {self.file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def extract_full_docx_content(self) -> str:
        '''Extrae texto completo del documento DOCX'''
        if not DOCX_AVAILABLE:
            logger.warning("python-docx no disponible para extracción de contenido")
            return ""
        
        try:
            doc = Document(self.file_path)
            full_text = []
            
            # Todos los párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Contenido de tablas
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    full_text.append("=== TABLA ===")
                    full_text.extend(table_text)
                    full_text.append("=== FIN TABLA ===")
            
            # Headers y footers de todas las secciones
            try:
                for section in doc.sections:
                    if hasattr(section, 'header') and section.header:
                        for header_para in section.header.paragraphs:
                            if header_para.text.strip():
                                full_text.append(f"=== HEADER === {header_para.text.strip()}")
                    
                    if hasattr(section, 'footer') and section.footer:
                        for footer_para in section.footer.paragraphs:
                            if footer_para.text.strip():
                                full_text.append(f"=== FOOTER === {footer_para.text.strip()}")
            except Exception as e:
                logger.warning(f"Error extrayendo headers/footers: {str(e)}")
            
            result = '\n'.join(full_text)
            logger.info(f"DOCX contenido extraído: {len(result)} caracteres")
            return result
            
        except Exception as e:
            logger.error(f"Error extrayendo contenido completo DOCX: {str(e)}")
            return ""
    
    def extract_full_doc_content(self) -> str:
        '''Extrae contenido de archivos DOC legacy'''
        try:
            # Método 1: docx2txt (más ligero)
            try:
                import docx2txt
                content = docx2txt.process(self.file_path)
                if content and content.strip():
                    logger.info(f"DOC contenido extraído con docx2txt: {len(content)} caracteres")
                    return content.strip()
            except ImportError:
                logger.warning("docx2txt no disponible")
            except Exception as e:
                logger.warning(f"Error con docx2txt: {str(e)}")
            
            # Método 2: textract (más robusto pero más dependencias)
            try:
                import textract
                content_bytes = textract.process(self.file_path)
                content = content_bytes.decode('utf-8', errors='ignore')
                if content and content.strip():
                    logger.info(f"DOC contenido extraído con textract: {len(content)} caracteres")
                    return content.strip()
            except ImportError:
                logger.warning("textract no disponible")
            except Exception as e:
                logger.warning(f"Error con textract: {str(e)}")
            
            logger.warning("No hay librerías disponibles para extraer contenido de DOC legacy")
            return ""
            
        except Exception as e:
            logger.error(f"Error extrayendo contenido DOC: {str(e)}")
            return ""
    
    def extract_full_xlsx_content(self) -> str:
        '''Extrae contenido completo de archivo XLSX'''
        if not OPENPYXL_AVAILABLE:
            logger.warning("openpyxl no disponible para extracción de contenido XLSX")
            return ""
        
        try:
            from django.conf import settings
            
            # Configuración de límites
            config = getattr(settings, 'CONTENT_EXTRACTION_SETTINGS', {})
            max_cells_per_sheet = config.get('MAX_EXCEL_CELLS_PER_SHEET', 1000)
            max_sheets = config.get('MAX_EXCEL_SHEETS', 10)
            
            workbook = load_workbook(self.file_path, data_only=True, read_only=True)
            sheets_content = []
            
            # Procesar hojas limitadas
            sheet_names = workbook.sheetnames[:max_sheets]
            logger.info(f"Procesando {len(sheet_names)} hojas de XLSX: {sheet_names}")
            
            for i, sheet_name in enumerate(sheet_names):
                sheet = workbook[sheet_name]
                sheet_content = self._extract_xlsx_sheet_content(
                    sheet, sheet_name, max_cells_per_sheet
                )
                if sheet_content:
                    sheets_content.append(sheet_content)
            
            # Convertir a texto plano para almacenamiento
            full_content = "\n\n".join(sheets_content)
            
            # Cerrar workbook para liberar memoria
            workbook.close()
            
            logger.info(f"XLSX contenido extraído: {len(full_content)} caracteres de {len(sheets_content)} hojas")
            return full_content
            
        except Exception as e:
            logger.error(f"Error extrayendo contenido completo XLSX: {str(e)}")
            return ""
    
    def _extract_xlsx_sheet_content(self, sheet, sheet_name: str, max_cells: int) -> str:
        '''Extrae contenido de una hoja XLSX específica'''
        try:
            content_lines = [f"=== HOJA: {sheet_name} ==="]
            cells_processed = 0
            
            # Verificar si la hoja tiene datos
            if sheet.max_row == 1 and sheet.max_column == 1:
                cell_value = sheet.cell(1, 1).value
                if cell_value is None:
                    content_lines.append("(Hoja vacía)")
                    return "\n".join(content_lines)
            
            # Procesar filas con datos (limitadas)
            max_rows = min(sheet.max_row, 100)
            max_cols = min(sheet.max_column, 50)
            
            for row in sheet.iter_rows(min_row=1, max_row=max_rows, 
                                      min_col=1, max_col=max_cols):
                if cells_processed >= max_cells:
                    content_lines.append(f"... (limitado a {max_cells} celdas)")
                    break
                    
                row_data = []
                has_data = False
                
                for cell in row:
                    if cell.value is not None:
                        has_data = True
                        cell_content = self._format_xlsx_cell_value(cell)
                        row_data.append(cell_content)
                        cells_processed += 1
                    else:
                        row_data.append("")
                
                if has_data and any(data.strip() for data in row_data):
                    content_lines.append(" | ".join(row_data))
            
            # Detectar nombres de rangos definidos
            try:
                if hasattr(sheet.parent, 'defined_names') and sheet.parent.defined_names:
                    range_names = []
                    for defined_name in sheet.parent.defined_names:
                        if defined_name.name and sheet_name in str(defined_name.value):
                            range_names.append(defined_name.name)
                    
                    if range_names:
                        content_lines.append("=== RANGOS DEFINIDOS ===")
                        content_lines.append(", ".join(range_names))
            except Exception as e:
                logger.warning(f"Error detectando rangos definidos: {str(e)}")
            
            return "\n".join(content_lines)
            
        except Exception as e:
            logger.warning(f"Error extrayendo hoja XLSX {sheet_name}: {str(e)}")
            return f"=== HOJA: {sheet_name} === (Error en extracción)"
    
    def _format_xlsx_cell_value(self, cell) -> str:
        '''Formatea el valor de una celda XLSX para extracción de texto'''
        try:
            if cell.value is None:
                return ""
            
            # Detectar fórmulas (si están disponibles)
            cell_formula = None
            try:
                if hasattr(cell, 'formula') and cell.formula:
                    cell_formula = cell.formula
            except:
                pass  # Algunas versiones no soportan fórmulas en read_only
            
            # Formatear según tipo de valor
            if isinstance(cell.value, (int, float)):
                value_str = str(cell.value)
                if cell_formula:
                    return f"{value_str} [={cell_formula}]"
                return value_str
            elif isinstance(cell.value, str):
                # Detectar patrones sensibles en contenido
                content = cell.value.strip()
                if self._contains_sensitive_pattern(content):
                    return f"[SENSITIVE]{content}"
                return content
            else:
                return str(cell.value)
                
        except Exception as e:
            return f"[ERROR:{str(e)}]"
    
    def _contains_sensitive_pattern(self, text: str) -> bool:
        '''Detecta patrones sensibles en texto de celda'''
        if not text or len(text.strip()) == 0:
            return False
            
        import re
        
        sensitive_patterns = [
            r'\b\d{1,2}\.\d{3}\.\d{3}-[\dkK]\b',  # RUT chileno
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Email
            r'\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b',  # Tarjeta/cuenta
            r'\b(password|passwd|clave|token|key)',  # Palabras clave sensibles
            r'\b(\+56|56)?\s?[9]\s?\d{4}\s?\d{4}\b',  # Teléfonos chilenos
        ]
        
        for pattern in sensitive_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def extract_full_xls_content(self) -> str:
        '''Extrae contenido completo de archivo XLS legacy'''
        try:
            # Intentar con xlrd primero
            try:
                import xlrd
                return self._extract_xls_with_xlrd()
            except ImportError:
                logger.warning("xlrd no disponible para archivos XLS legacy")
            except Exception as e:
                logger.warning(f"Error con xlrd: {str(e)}")
            
            # Fallback: intentar con openpyxl (puede leer algunos XLS modernos)
            if OPENPYXL_AVAILABLE:
                logger.info("Intentando XLS con openpyxl como fallback")
                return self.extract_full_xlsx_content()  # Reutilizar método XLSX
            
            logger.warning("No hay librerías disponibles para extraer contenido XLS")
            return ""
            
        except Exception as e:
            logger.error(f"Error extrayendo contenido completo XLS: {str(e)}")
            return ""
    
    def _extract_xls_with_xlrd(self) -> str:
        '''Extrae contenido XLS usando xlrd'''
        import xlrd
        from django.conf import settings
        
        # Configuración de límites
        config = getattr(settings, 'CONTENT_EXTRACTION_SETTINGS', {})
        max_cells_per_sheet = config.get('MAX_EXCEL_CELLS_PER_SHEET', 1000)
        max_sheets = config.get('MAX_EXCEL_SHEETS', 10)
        
        try:
            workbook = xlrd.open_workbook(self.file_path)
            sheets_content = []
            
            # Procesar hojas limitadas
            sheet_names = workbook.sheet_names()[:max_sheets]
            logger.info(f"Procesando {len(sheet_names)} hojas de XLS: {sheet_names}")
            
            for sheet_name in sheet_names:
                sheet = workbook.sheet_by_name(sheet_name)
                sheet_content = self._extract_xls_sheet_content(sheet, sheet_name, max_cells_per_sheet)
                if sheet_content:
                    sheets_content.append(sheet_content)
            
            full_content = "\n\n".join(sheets_content)
            logger.info(f"XLS contenido extraído: {len(full_content)} caracteres de {len(sheets_content)} hojas")
            return full_content
            
        except Exception as e:
            logger.error(f"Error procesando XLS con xlrd: {str(e)}")
            return ""
    
    def _extract_xls_sheet_content(self, sheet, sheet_name: str, max_cells: int) -> str:
        '''Extrae contenido de una hoja XLS específica'''
        try:
            content_lines = [f"=== HOJA: {sheet_name} ==="]
            cells_processed = 0
            
            # Verificar si la hoja tiene datos
            if sheet.nrows == 0 or sheet.ncols == 0:
                content_lines.append("(Hoja vacía)")
                return "\n".join(content_lines)
            
            # Procesar filas limitadas
            max_rows = min(sheet.nrows, 100)
            max_cols = min(sheet.ncols, 50)
            
            for row_idx in range(max_rows):
                if cells_processed >= max_cells:
                    content_lines.append(f"... (limitado a {max_cells} celdas)")
                    break
                    
                row_data = []
                has_data = False
                
                for col_idx in range(max_cols):
                    try:
                        cell_value = sheet.cell_value(row_idx, col_idx)
                        if cell_value is not None and cell_value != "":
                            has_data = True
                            formatted_value = self._format_xls_cell_value(cell_value)
                            row_data.append(formatted_value)
                            cells_processed += 1
                        else:
                            row_data.append("")
                    except Exception as e:
                        row_data.append(f"[ERROR_CELL]")
                        logger.warning(f"Error leyendo celda [{row_idx},{col_idx}]: {str(e)}")
                
                if has_data and any(data.strip() for data in row_data):
                    content_lines.append(" | ".join(row_data))
            
            return "\n".join(content_lines)
            
        except Exception as e:
            logger.warning(f"Error extrayendo hoja XLS {sheet_name}: {str(e)}")
            return f"=== HOJA: {sheet_name} === (Error en extracción)"
    
    def _format_xls_cell_value(self, cell_value) -> str:
        '''Formatea valor de celda XLS'''
        try:
            if cell_value is None or cell_value == "":
                return ""
            
            if isinstance(cell_value, (int, float)):
                return str(cell_value)
            elif isinstance(cell_value, str):
                content = cell_value.strip()
                if self._contains_sensitive_pattern(content):
                    return f"[SENSITIVE]{content}"
                return content
            else:
                return str(cell_value)
                
        except Exception as e:
            return f"[ERROR:{str(e)}]"
    
    def _extract_xlsx_metadata(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        '''Extrae metadatos de hojas de cálculo Excel (.xlsx)'''
        if not OPENPYXL_AVAILABLE:
            metadata['extraction_error'] = 'openpyxl no disponible'
            return metadata
        
        try:
            workbook = load_workbook(self.file_path, data_only=True)
            
            office_metadata = {}
            
            # Propiedades del workbook con validación defensiva
            props = workbook.properties
            
            # Lista de propiedades a extraer con validación
            property_mapping = [
                ('creator', 'creator'),
                ('lastModifiedBy', 'last_modified_by'),
                ('title', 'title'),
                ('subject', 'subject'),
                ('keywords', 'keywords'),
                ('description', 'description'),
                ('category', 'category'),
                ('company', 'company'),
            ]
            
            for prop_name, meta_name in property_mapping:
                if hasattr(props, prop_name):
                    prop_value = getattr(props, prop_name)
                    if prop_value:  # Solo agregar si no está vacío
                        office_metadata[meta_name] = str(prop_value)
            
            # Fechas con manejo especial
            if hasattr(props, 'created') and props.created:
                office_metadata['created'] = props.created.isoformat()
            if hasattr(props, 'modified') and props.modified:
                office_metadata['modified'] = props.modified.isoformat()
            
            metadata['office_metadata'] = office_metadata
            
            # Información de las hojas
            sheets_info = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheets_info.append({
                    'name': sheet_name,
                    'max_row': sheet.max_row,
                    'max_column': sheet.max_column
                })
            
            metadata['sheets_info'] = sheets_info
            metadata['document_info'] = {
                'sheets_count': len(workbook.sheetnames),
                'document_type': 'Excel Workbook (.xlsx)'
            }
            
        except Exception as e:
            logger.error(f"Error extrayendo metadatos de XLSX {self.file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _extract_xls_metadata(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        '''Extrae metadatos de hojas de cálculo Excel legacy (.xls)'''
        try:
            # Intentar con xlrd primero (para .xls legacy)
            try:
                import xlrd
                
                workbook = xlrd.open_workbook(self.file_path)
                
                office_metadata = {
                    'document_type': 'Excel Workbook Legacy (.xls)',
                    'sheets_count': workbook.nsheets,
                }
                
                # Información de hojas
                sheets_info = []
                for sheet_name in workbook.sheet_names():
                    sheet = workbook.sheet_by_name(sheet_name)
                    sheets_info.append({
                        'name': sheet_name,
                        'nrows': sheet.nrows,
                        'ncols': sheet.ncols
                    })
                
                metadata['sheets_info'] = sheets_info
                metadata['office_metadata'] = office_metadata
                
                return metadata
                
            except ImportError:
                logger.warning("xlrd no disponible, intentando con openpyxl para archivo .xls")
                # Fallback: intentar con openpyxl (puede manejar algunos .xls)
                if OPENPYXL_AVAILABLE:
                    return self._extract_xlsx_metadata(metadata)
                else:
                    metadata['extraction_error'] = 'Ni xlrd ni openpyxl disponibles para archivos .xls'
                    return metadata
            
        except Exception as e:
            logger.error(f"Error extrayendo metadatos de XLS {self.file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
        
        return metadata


class OpenOfficeExtractor(MetadataExtractor):
    '''Extractor especializado para archivos OpenDocument Format (.odt, .ods, .odp)'''
    
    def extract(self) -> Dict[str, Any]:
        metadata = self.get_common_metadata()
        file_extension = os.path.splitext(self.file_path)[1].lower()
        
        try:
            if ODFPY_AVAILABLE:
                # Usar odfpy para extracción completa
                odf_metadata = self._extract_with_odfpy(file_extension)
            else:
                # Fallback a método ZIP directo
                odf_metadata = self._extract_with_zip_fallback(file_extension)
                
            metadata.update(odf_metadata)
            
        except Exception as e:
            logger.error(f"Error extrayendo metadatos ODF {self.file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
            # IMPORTANTE: No interrumpir el crawling, continuar con metadatos básicos
            
        return metadata
    
    def extract_full_odt_content(self) -> str:
        '''Extrae texto completo del documento ODT'''
        try:
            if not ODFPY_AVAILABLE:
                logger.warning("odfpy no disponible para extracción de contenido ODT")
                return ""
            
            from odf.opendocument import load
            from odf import text, table
            
            doc = load(self.file_path)
            full_text = []
            
            # Obtener todos los elementos de texto
            all_text_elements = []
            all_text_elements.extend(doc.getElementsByType(text.H))  # Headers
            all_text_elements.extend(doc.getElementsByType(text.P))  # Paragraphs
            
            for element in all_text_elements:
                element_text = self._extract_element_text(element)
                if element_text.strip():
                    full_text.append(element_text.strip())
            
            # Texto de tablas
            tables = doc.getElementsByType(table.Table)
            for table_elem in tables:
                table_text = self._extract_full_table_text(table_elem)
                if table_text:
                    full_text.append("=== TABLA ===")
                    full_text.append(table_text)
                    full_text.append("=== FIN TABLA ===")
            
            result = '\n'.join(full_text)
            logger.info(f"ODT contenido extraído: {len(result)} caracteres")
            return result
            
        except Exception as e:
            logger.error(f"Error extrayendo contenido completo ODT: {str(e)}")
            return ""
    
    def _extract_full_table_text(self, table_elem) -> str:
        '''Extrae todo el texto de una tabla ODT'''
        try:
            from odf import table
            table_text = []
            
            rows = table_elem.getElementsByType(table.TableRow)
            for row in rows:
                row_text = []
                cells = row.getElementsByType(table.TableCell)
                for cell in cells:
                    cell_text = self._extract_element_text(cell)
                    if cell_text.strip():
                        row_text.append(cell_text.strip())
                if row_text:
                    table_text.append(' | '.join(row_text))
            
            return '\n'.join(table_text)
            
        except Exception as e:
            logger.warning(f"Error extrayendo tabla ODT: {str(e)}")
            return ""
    
    def _extract_with_odfpy(self, file_extension: str) -> Dict[str, Any]:
        '''Extrae metadatos usando la librería odfpy'''
        metadata = {}
        
        try:
            # Cargar documento ODF
            doc = load(self.file_path)
            
            # Extraer metadatos comunes
            doc_meta = self._extract_document_metadata(doc)
            metadata.update(doc_meta)
            
            # Extraer metadatos específicos por tipo
            if file_extension == '.odt':
                specific_meta = self._extract_odt_specific(doc)
            elif file_extension == '.ods':
                specific_meta = self._extract_ods_specific(doc)
            elif file_extension == '.odp':
                specific_meta = self._extract_odp_specific(doc)
            else:
                specific_meta = {'document_type': f'OpenDocument {file_extension}'}
                
            metadata.update(specific_meta)
            
            logger.info(f"Metadatos ODF extraídos exitosamente de {self.file_path}")
            
        except Exception as e:
            logger.error(f"Error con odfpy para {self.file_path}: {str(e)}")
            # Intentar fallback
            return self._extract_with_zip_fallback(file_extension)
            
        return metadata
    
    def _extract_document_metadata(self, doc) -> Dict[str, Any]:
        '''Extrae metadatos comunes del documento ODF'''
        openoffice_metadata = {}
        
        try:
            from odf import meta as odf_meta, dc
            
            meta_element = doc.meta
            
            # Lista de elementos meta a extraer con sus tipos ODF correspondientes
            meta_extractors = [
                (dc.Title, 'title'),
                (dc.Subject, 'subject'),
                (dc.Description, 'description'), 
                (dc.Creator, 'creator'),
                (odf_meta.InitialCreator, 'initial_creator'),
                (odf_meta.CreationDate, 'creation_date'),
                (dc.Date, 'modification_date'),
                (dc.Language, 'language'),
                (odf_meta.Keyword, 'keywords'),
                (odf_meta.Generator, 'generator'),
                (odf_meta.EditingCycles, 'editing_cycles'),
                (odf_meta.EditingDuration, 'editing_duration'),
                (odf_meta.Template, 'template')
            ]
            
            for element_type, field_name in meta_extractors:
                try:
                    elements = meta_element.getElementsByType(element_type)
                    if elements and len(elements) > 0:
                        # Extraer texto del primer elemento encontrado
                        element_text = self._get_element_text_content(elements[0])
                        if element_text and element_text.strip():
                            openoffice_metadata[field_name] = element_text.strip()
                except Exception:
                    # Si falla la extracción de un campo específico, continuar con el siguiente
                    continue
            
            if openoffice_metadata:
                return {'openoffice_metadata': openoffice_metadata}
            else:
                return {}
            
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos básicos ODF: {str(e)}")
            return {}
    
    def _get_element_text_content(self, element) -> str:
        '''Obtiene el contenido de texto de un elemento ODF'''
        try:
            # Intentar obtener texto del elemento y sus hijos
            text_parts = []
            
            # Texto directo del elemento
            if hasattr(element, 'data') and element.data:
                text_parts.append(str(element.data))
            
            # Texto de nodos hijos
            if hasattr(element, 'childNodes'):
                for child in element.childNodes:
                    if hasattr(child, 'data') and child.data:
                        text_parts.append(str(child.data))
                    elif hasattr(child, 'childNodes'):
                        child_text = self._get_element_text_content(child)
                        if child_text:
                            text_parts.append(child_text)
            
            return ' '.join(text_parts).strip()
            
        except Exception:
            return ''
    
    def _extract_odt_specific(self, doc) -> Dict[str, Any]:
        '''Extrae metadatos específicos de documentos de texto ODT'''
        try:
            # Información del documento
            document_info = {
                'document_type': 'OpenDocument Text (.odt)'
            }
            
            # Contar elementos del documento
            paragraphs = doc.getElementsByType(text.P)
            headers = doc.getElementsByType(text.H)
            images = doc.getElementsByType(draw.Image)
            
            document_info['paragraph_count'] = len(paragraphs)
            document_info['header_count'] = len(headers)
            document_info['image_count'] = len(images)
            
            # Extraer vista previa del contenido (primeros párrafos)
            content_preview = self._extract_text_preview(doc)
            
            result = {
                'document_info': document_info
            }
            
            if content_preview:
                result['content_preview'] = content_preview
                
            return result
            
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos ODT específicos: {str(e)}")
            return {'document_info': {'document_type': 'OpenDocument Text (.odt)'}}
    
    def _extract_ods_specific(self, doc) -> Dict[str, Any]:
        '''Extrae metadatos específicos de hojas de cálculo ODS'''
        try:
            # Información del documento
            document_info = {
                'document_type': 'OpenDocument Spreadsheet (.ods)'
            }
            
            # Obtener todas las hojas
            spreadsheets = doc.getElementsByType(table.Table)
            document_info['sheet_count'] = len(spreadsheets)
            
            # Información de cada hoja
            sheets_info = []
            for i, sheet in enumerate(spreadsheets[:5]):  # Máximo 5 hojas
                sheet_info = {
                    'name': sheet.getAttribute('name') or f'Sheet{i+1}',
                }
                
                # Contar filas y columnas con datos
                rows = sheet.getElementsByType(table.TableRow)
                if rows:
                    sheet_info['row_count'] = len(rows)
                    
                sheets_info.append(sheet_info)
            
            result = {
                'document_info': document_info,
                'sheets_info': sheets_info
            }
            
            # Vista previa de contenido de la primera hoja
            if sheets_info:
                content_preview = self._extract_spreadsheet_preview(spreadsheets[0])
                if content_preview:
                    result['content_preview'] = content_preview
            
            return result
            
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos ODS específicos: {str(e)}")
            return {'document_info': {'document_type': 'OpenDocument Spreadsheet (.ods)'}}
    
    def _extract_odp_specific(self, doc) -> Dict[str, Any]:
        '''Extrae metadatos específicos de presentaciones ODP'''
        try:
            # Información del documento
            document_info = {
                'document_type': 'OpenDocument Presentation (.odp)'
            }
            
            # Contar slides (páginas de la presentación)
            pages = doc.getElementsByType(draw.Page)
            document_info['slide_count'] = len(pages)
            
            # Vista previa de los primeros slides
            slides_preview = []
            for i, page in enumerate(pages[:3]):  # Primeras 3 slides
                slide_info = {
                    'slide_number': i + 1,
                    'slide_name': page.getAttribute('name') or f'Slide {i+1}'
                }
                
                # Extraer texto de la slide si es posible
                slide_text = self._extract_slide_text(page)
                if slide_text:
                    slide_info['text_preview'] = slide_text[:200]
                    
                slides_preview.append(slide_info)
            
            result = {
                'document_info': document_info
            }
            
            if slides_preview:
                result['slides_preview'] = slides_preview
                # Usar el texto de las slides como preview general
                all_text = ' '.join([slide.get('text_preview', '') for slide in slides_preview])
                if all_text.strip():
                    result['content_preview'] = all_text[:500]
            
            return result
            
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos ODP específicos: {str(e)}")
            return {'document_info': {'document_type': 'OpenDocument Presentation (.odp)'}}
    
    def _extract_text_preview(self, doc) -> str:
        '''Extrae vista previa de texto de un documento ODT'''
        try:
            # Obtener párrafos y headers
            paragraphs = doc.getElementsByType(text.P)
            headers = doc.getElementsByType(text.H)
            
            # Combinar todo el texto
            all_elements = headers + paragraphs
            
            text_content = []
            for element in all_elements[:10]:  # Primeros 10 elementos
                element_text = self._extract_element_text(element)
                if element_text.strip():
                    text_content.append(element_text.strip())
                
            preview = ' '.join(text_content)
            return preview[:500] if preview else ''
            
        except Exception as e:
            logger.warning(f"Error extrayendo preview de texto: {str(e)}")
            return ''
    
    def _extract_spreadsheet_preview(self, sheet) -> str:
        '''Extrae vista previa de una hoja de cálculo'''
        try:
            rows = sheet.getElementsByType(table.TableRow)
            preview_lines = []
            
            for i, row in enumerate(rows[:5]):  # Primeras 5 filas
                cells = row.getElementsByType(table.TableCell)
                row_data = []
                
                for j, cell in enumerate(cells[:5]):  # Primeras 5 columnas
                    cell_text = self._extract_element_text(cell)
                    if cell_text.strip():
                        row_data.append(cell_text.strip())
                        
                if row_data:
                    preview_lines.append(' | '.join(row_data))
                    
            return '\n'.join(preview_lines)[:500] if preview_lines else ''
            
        except Exception as e:
            logger.warning(f"Error extrayendo preview de hoja: {str(e)}")
            return ''
    
    def _extract_slide_text(self, slide) -> str:
        '''Extrae texto de una slide de presentación'''
        try:
            # Buscar elementos de texto en la slide
            text_boxes = slide.getElementsByType(draw.TextBox)
            slide_text = []
            
            for text_box in text_boxes:
                box_text = self._extract_element_text(text_box)
                if box_text.strip():
                    slide_text.append(box_text.strip())
            
            return ' '.join(slide_text)
            
        except Exception as e:
            logger.warning(f"Error extrayendo texto de slide: {str(e)}")
            return ''
    
    def _extract_element_text(self, element) -> str:
        '''Extrae texto de un elemento ODF de manera recursiva'''
        try:
            text_content = []
            
            # Obtener texto directo del elemento
            if hasattr(element, 'data'):
                text_content.append(str(element.data))
            
            # Obtener texto de elementos hijos
            for child in element.childNodes:
                if hasattr(child, 'data'):
                    text_content.append(str(child.data))
                elif hasattr(child, 'childNodes'):
                    # Recursión para elementos anidados
                    child_text = self._extract_element_text(child)
                    if child_text:
                        text_content.append(child_text)
            
            return ' '.join(text_content)
            
        except Exception:
            return ''
    
    def _extract_with_zip_fallback(self, file_extension: str) -> Dict[str, Any]:
        '''Método de fallback usando ZIP directo para extraer metadatos básicos'''
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            logger.info(f"Usando método ZIP fallback para {self.file_path}")
            
            metadata = {}
            
            with zipfile.ZipFile(self.file_path, 'r') as zip_file:
                # Intentar leer meta.xml
                try:
                    with zip_file.open('meta.xml') as meta_file:
                        meta_xml = meta_file.read()
                        root = ET.fromstring(meta_xml)
                        
                        # Extraer metadatos básicos del XML
                        basic_meta = self._parse_meta_xml(root)
                        if basic_meta:
                            metadata['openoffice_metadata'] = basic_meta
                            
                except (KeyError, ET.ParseError):
                    logger.warning("No se pudo procesar meta.xml")
                
                # Información básica del documento
                metadata['document_info'] = {
                    'document_type': f'OpenDocument {file_extension}',
                    'extraction_method': 'ZIP fallback'
                }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error en método ZIP fallback: {str(e)}")
            return {
                'document_info': {
                    'document_type': f'OpenDocument {file_extension}',
                    'extraction_method': 'fallback failed'
                },
                'extraction_error': str(e)
            }
    
    def _parse_meta_xml(self, root) -> Dict[str, Any]:
        '''Parsea el archivo meta.xml para extraer metadatos básicos'''
        try:
            metadata = {}
            
            # Namespace común de OpenDocument
            namespaces = {
                'meta': 'urn:oasis:names:tc:opendocument:xmlns:meta:1.0',
                'dc': 'http://purl.org/dc/elements/1.1/'
            }
            
            # Mapeo de campos XML a nuestros campos
            field_mapping = [
                ('.//dc:title', 'title'),
                ('.//dc:subject', 'subject'),
                ('.//dc:description', 'description'),
                ('.//dc:creator', 'creator'),
                ('.//dc:date', 'modification_date'),
                ('.//meta:creation-date', 'creation_date'),
                ('.//meta:generator', 'generator'),
                ('.//meta:keyword', 'keywords')
            ]
            
            for xpath, field_name in field_mapping:
                try:
                    element = root.find(xpath, namespaces)
                    if element is not None and element.text:
                        metadata[field_name] = element.text.strip()
                except Exception:
                    continue
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error parseando meta.xml: {str(e)}")
            return {}


class MultimediaExtractor(MetadataExtractor):
    '''Extractor para archivos multimedia (MP3, MP4, etc.)'''
    
    def extract(self) -> Dict[str, Any]:
        metadata = self.get_common_metadata()
        
        if not MUTAGEN_AVAILABLE:
            metadata['extraction_error'] = 'mutagen no disponible'
            return metadata
        
        try:
            audio_file = MutagenFile(self.file_path)
            
            if audio_file is None:
                metadata['extraction_error'] = 'Formato multimedia no soportado'
                return metadata
            
            multimedia_metadata = {}
            
            # Metadatos comunes de audio
            if hasattr(audio_file, 'tags') and audio_file.tags:
                tags = audio_file.tags
                
                # MP3 tags
                if isinstance(audio_file, MP3):
                    multimedia_metadata.update(self._extract_mp3_tags(tags))
                # MP4 tags
                elif isinstance(audio_file, MP4):
                    multimedia_metadata.update(self._extract_mp4_tags(tags))
                else:
                    # Tags genéricos
                    for key, value in tags.items():
                        multimedia_metadata[str(key)] = str(value[0]) if isinstance(value, list) else str(value)
            
            # Información del archivo multimedia
            if hasattr(audio_file, 'info'):
                info = audio_file.info
                multimedia_metadata['duration_seconds'] = getattr(info, 'length', 0)
                multimedia_metadata['bitrate'] = getattr(info, 'bitrate', 0)
                multimedia_metadata['sample_rate'] = getattr(info, 'sample_rate', 0)
                multimedia_metadata['channels'] = getattr(info, 'channels', 0)
            
            metadata['multimedia_metadata'] = multimedia_metadata
            
        except Exception as e:
            logger.error(f"Error extrayendo metadatos multimedia {self.file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _extract_mp3_tags(self, tags) -> Dict[str, Any]:
        '''Extrae tags específicos de MP3'''
        mp3_metadata = {}
        
        # Mapeo de tags ID3
        tag_mapping = {
            'TIT2': 'title',
            'TPE1': 'artist',
            'TALB': 'album',
            'TDRC': 'year',
            'TCON': 'genre',
            'TRCK': 'track_number',
            'TPE2': 'album_artist',
            'TPOS': 'disc_number',
            'TCOP': 'copyright',
            'TENC': 'encoded_by',
            'TSSE': 'encoding_software',
        }
        
        for tag_id, tag_name in tag_mapping.items():
            if tag_id in tags:
                mp3_metadata[tag_name] = str(tags[tag_id][0])
        
        return mp3_metadata
    
    def _extract_mp4_tags(self, tags) -> Dict[str, Any]:
        '''Extrae tags específicos de MP4'''
        mp4_metadata = {}
        
        # Mapeo de tags MP4
        tag_mapping = {
            '\xa9nam': 'title',
            '\xa9ART': 'artist',
            '\xa9alb': 'album',
            '\xa9day': 'year',
            '\xa9gen': 'genre',
            'trkn': 'track_number',
            '\xa9cpy': 'copyright',
            '\xa9too': 'encoding_software',
        }
        
        for tag_id, tag_name in tag_mapping.items():
            if tag_id in tags:
                value = tags[tag_id][0]
                mp4_metadata[tag_name] = str(value)
        
        return mp4_metadata


class HTMLExtractor(MetadataExtractor):
    '''Extractor para páginas HTML'''
    
    def extract(self) -> Dict[str, Any]:
        metadata = self.get_common_metadata()
        
        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            
            html_metadata = {}
            
            # Meta tags
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                if meta.get('name'):
                    html_metadata[f"meta_{meta.get('name')}"] = meta.get('content', '')
                elif meta.get('property'):
                    html_metadata[f"property_{meta.get('property')}"] = meta.get('content', '')
            
            # Título
            title_tag = soup.find('title')
            if title_tag:
                html_metadata['title'] = title_tag.get_text().strip()
            
            # Enlaces
            links = soup.find_all('a', href=True)
            html_metadata['links_count'] = len(links)
            
            # Imágenes
            images = soup.find_all('img', src=True)
            html_metadata['images_count'] = len(images)
            
            metadata['html_metadata'] = html_metadata
            
        except Exception as e:
            logger.error(f"Error extrayendo metadatos HTML {self.file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
        
        return metadata


# Factory function para obtener el extractor apropiado
def get_metadata_extractor(file_path: str, file_url: str = None, referrer: str = None) -> MetadataExtractor:
    '''
    Factory function que retorna el extractor apropiado según el tipo de archivo
    '''
    file_extension = os.path.splitext(file_path)[1].lower()
    
    # Mapeo de extensiones a extractores
    if file_extension == '.pdf':
        return PDFExtractor(file_path, file_url, referrer)
    elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.gif']:
        return ImageExtractor(file_path, file_url, referrer)
    elif file_extension in ['.docx', '.xlsx', '.pptx', '.xls', '.doc', '.ppt']:
        return OfficeExtractor(file_path, file_url, referrer)
    elif file_extension in ['.odt', '.ods', '.odp']:
        return OpenOfficeExtractor(file_path, file_url, referrer)
    elif file_extension in ['.mp3', '.mp4']:
        return MultimediaExtractor(file_path, file_url, referrer)
    elif file_extension in ['.html', '.htm']:
        return HTMLExtractor(file_path, file_url, referrer)
    else:
        # Extractor genérico para otros tipos
        return MetadataExtractor(file_path, file_url, referrer)


def extract_metadata_from_file(file_path: str, file_url: str = None, referrer: str = None) -> Dict[str, Any]:
    '''
    Función principal para extraer metadatos de cualquier archivo
    '''
    try:
        extractor = get_metadata_extractor(file_path, file_url, referrer)
        return extractor.extract()
    except Exception as e:
        logger.error(f"Error en extracción de metadatos para {file_path}: {str(e)}")
        return {
            'file_path': file_path,
            'file_url': file_url,
            'referrer': referrer,
            'extraction_error': str(e),
            'extracted_at': datetime.now().isoformat(),
        }


def extract_full_content_from_file(file_path: str) -> str:
    '''Extrae contenido completo de un archivo según su tipo'''
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            extractor = PDFExtractor(file_path)
            return extractor.extract_full_content()
        elif file_extension == '.docx':
            extractor = OfficeExtractor(file_path)
            return extractor.extract_full_docx_content()
        elif file_extension == '.doc':
            extractor = OfficeExtractor(file_path)
            return extractor.extract_full_doc_content()
        elif file_extension == '.odt':
            extractor = OpenOfficeExtractor(file_path)
            return extractor.extract_full_odt_content()
        # NUEVOS: Soporte para Excel
        elif file_extension == '.xlsx':
            extractor = OfficeExtractor(file_path)
            return extractor.extract_full_xlsx_content()
        elif file_extension == '.xls':
            extractor = OfficeExtractor(file_path)
            return extractor.extract_full_xls_content()
        else:
            logger.info(f"Tipo de archivo no soportado para extracción completa: {file_extension}")
            return ""
            
    except Exception as e:
        logger.error(f"Error extrayendo contenido completo de {file_path}: {str(e)}")
        return ""